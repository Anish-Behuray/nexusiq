"""
ingestion/loaders/document_loader.py
======================================
Document loaders: convert raw files into LlamaIndex Document objects.

CONCEPT: What is a "Document" in the RAG context?
----------------------------------------------------
A LlamaIndex Document is a standardized wrapper around raw text.
It contains:
  - text: the actual content
  - metadata: information ABOUT the content
    (filename, page number, author, date, source URL)

Metadata is crucial because it:
1. Enables citations: "This answer came from page 3 of policy.pdf"
2. Enables filtering: "Only search HR documents"
3. Enables re-ingestion: "Re-process only modified files"

SUPPORTED FORMATS: PDF, Markdown, Plain text, Word (.docx)

INDUSTRY PRACTICE: The loader layer is "dumb" — it just converts
raw bytes to text + metadata. All the intelligence (chunking,
embedding, retrieval) happens in later pipeline stages.
This separation makes it easy to add new file types.
"""

import os
from pathlib import Path
from typing import Optional
from loguru import logger

from llama_index.core import Document
from llama_index.core.node_parser import SimpleNodeParser


def load_document(file_path: str, document_id: str) -> list[Document]:
    """
    Load a file and return a list of LlamaIndex Documents.
    
    Why a list? Some formats (like PDF) are better loaded
    page-by-page, so one PDF → multiple Document objects
    (one per page). This gives us precise page citations.
    
    Args:
        file_path: Absolute path to the file on disk
        document_id: Our internal DB document ID (for linking)
    
    Returns:
        List of Document objects ready for chunking
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    
    logger.info(f"Loading document: {path.name} (type={ext})")
    
    loaders = {
        ".pdf": _load_pdf,
        ".md": _load_markdown,
        ".txt": _load_text,
        ".docx": _load_docx,
    }
    
    loader_fn = loaders.get(ext)
    if not loader_fn:
        raise ValueError(f"Unsupported file type: {ext}")
    
    documents = loader_fn(file_path, document_id)
    
    logger.info(f"Loaded {len(documents)} document sections from {path.name}")
    return documents


def _base_metadata(file_path: str, document_id: str) -> dict:
    """
    Build standard metadata that ALL loaders include.
    
    Consistent metadata across all document types means our
    retrieval layer can always access filename, doc_id, etc.
    """
    path = Path(file_path)
    return {
        "document_id": document_id,       # Links to PostgreSQL
        "filename": path.name,
        "file_type": path.suffix.lstrip("."),
        "file_size_bytes": os.path.getsize(file_path),
    }


def _load_pdf(file_path: str, document_id: str) -> list[Document]:
    """
    Load PDF, splitting by page for precise citations.
    
    We use pypdf for reliability and speed.
    Each page becomes its own Document with page_number metadata.
    
    WHY PAGE-BY-PAGE:
    If we loaded the whole PDF as one document, we'd lose page info.
    Citations like "page 7" are essential for enterprise users who
    need to verify AI answers in the original document.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Install pypdf: pip install pypdf")
    
    base_meta = _base_metadata(file_path, document_id)
    documents = []
    
    reader = PdfReader(file_path)
    
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        
        if not text or not text.strip():
            # Skip blank pages (common in PDFs with images/diagrams)
            logger.debug(f"Skipping blank page {page_num}")
            continue
        
        doc = Document(
            text=text,
            metadata={
                **base_meta,
                "page_number": page_num,
                "total_pages": len(reader.pages),
                # Section detection (very basic — Phase 2 uses layout-aware models)
                "section_header": _extract_first_heading(text),
            }
        )
        documents.append(doc)
    
    if not documents:
        raise ValueError(f"PDF has no extractable text: {file_path}")
    
    return documents


def _load_markdown(file_path: str, document_id: str) -> list[Document]:
    """
    Load Markdown, splitting by H2/H3 headers.
    
    Markdown documents (Notion exports, READMEs, wikis) are
    naturally structured with headers. Splitting on headers
    means each chunk is a self-contained section — better
    retrieval quality than arbitrary fixed-size chunks.
    """
    base_meta = _base_metadata(file_path, document_id)
    
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    
    if not content.strip():
        raise ValueError(f"Markdown file is empty: {file_path}")
    
    # Split by H2 headers (## Section Title)
    sections = _split_by_markdown_headers(content)
    
    documents = []
    for section_idx, (header, text) in enumerate(sections):
        doc = Document(
            text=text,
            metadata={
                **base_meta,
                "section_header": header or "Introduction",
                "section_index": section_idx,
            }
        )
        documents.append(doc)
    
    return documents if documents else [
        # If no headers, return as single document
        Document(text=content, metadata={**base_meta, "section_header": "Full document"})
    ]


def _load_text(file_path: str, document_id: str) -> list[Document]:
    """
    Load plain text files — simplest case, one Document per file.
    """
    base_meta = _base_metadata(file_path, document_id)
    
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    
    if not content.strip():
        raise ValueError(f"Text file is empty: {file_path}")
    
    return [Document(text=content, metadata=base_meta)]


def _load_docx(file_path: str, document_id: str) -> list[Document]:
    """
    Load Word documents, preserving heading structure.
    
    python-docx lets us read paragraph styles (Heading 1, Heading 2)
    so we can split on section boundaries — same benefit as Markdown.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")
    
    base_meta = _base_metadata(file_path, document_id)
    doc = DocxDocument(file_path)
    
    # Collect paragraphs and split on headings
    sections = []
    current_header = "Introduction"
    current_text = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Detect heading paragraphs
        if paragraph.style.name.startswith("Heading"):
            # Save previous section
            if current_text:
                sections.append((current_header, "\n".join(current_text)))
            current_header = text
            current_text = []
        else:
            current_text.append(text)
    
    # Don't forget the last section
    if current_text:
        sections.append((current_header, "\n".join(current_text)))
    
    return [
        Document(
            text=text,
            metadata={**base_meta, "section_header": header, "section_index": i}
        )
        for i, (header, text) in enumerate(sections)
        if text.strip()  # Skip empty sections
    ]


# ─── Helper Functions ────────────────────────────────────────

def _extract_first_heading(text: str) -> Optional[str]:
    """
    Extract the first line that looks like a heading from plain text.
    Used for PDF pages that have implicit headings (all-caps lines, etc.)
    """
    lines = text.split("\n")
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if line and len(line) < 100:  # Short lines are likely headings
            return line
    return None


def _split_by_markdown_headers(content: str) -> list[tuple[str, str]]:
    """
    Split markdown content by H2 (##) headers.
    Returns list of (header_text, section_content) tuples.
    """
    import re
    
    # Split on lines starting with ## (H2 headers)
    sections = re.split(r'\n(?=## )', content)
    result = []
    
    for section in sections:
        lines = section.split("\n")
        if not lines:
            continue
        
        # First line is the header (if it starts with #)
        first_line = lines[0].strip()
        if first_line.startswith("#"):
            header = first_line.lstrip("#").strip()
            body = "\n".join(lines[1:]).strip()
        else:
            header = None
            body = section.strip()
        
        if body:  # Only include non-empty sections
            result.append((header, body))
    
    return result
